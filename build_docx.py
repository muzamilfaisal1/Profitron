"""Build a Wiley-style two-column research paper DOCX from research_paper.md.

Layout matches the Wiley HBET journal sample:
  - Page 1: single-column title block (article type, title, authors, affiliations,
    correspondence, dates, copyright, abstract paragraph, keywords).
  - From Section 1 onwards: two-column justified body with numbered sections
    (1. Introduction) and italicised subsection labels (1.1. Background.).
  - Times-style serif throughout, page numbers on outer header.
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


SRC = Path('research_paper.md')
OUT = Path('research_paper.docx')
FIG_DIR = Path('figures')

BODY_FONT = 'Times New Roman'
BODY_SIZE = Pt(9.5)
CAPTION_SIZE = Pt(8.5)
HEADING_SIZE = Pt(11)
SUBHEADING_SIZE = Pt(9.5)
TITLE_SIZE = Pt(20)


# Map figure file stem to its description (figure number is auto-assigned).
DESCRIPTIONS = {
    'fig01_architecture': 'Three-layer system architecture of the proposed trading agent.',
    'fig02_lstm_cell': 'Internal structure of a single LSTM cell used in the forecasting layer.',
    'fig03_training_curves': 'Training and validation dynamics of the LSTM forecasting model.',
    'fig04_prediction_vs_actual': 'Forecast vs. realised price on a held-out ETH/USDT sample.',
    'fig05_equity_curve': 'Backtested equity curve over the six-month evaluation window.',
    'fig06_regime_performance': 'Regime-conditional performance of the LSTM forecaster.',
    'fig07_latency_distribution': 'Distribution of observed order-fill latencies (n = 1,000).',
    'fig08_confusion_matrix': 'Confusion matrix of the three-way directional classifier.',
    'fig09_accuracy_by_asset': 'Per-asset directional accuracy of five forecasting models.',
    'fig10_data_pipeline': 'End-to-end data and control flow through the trading pipeline.',
    'fig11_sharpe_drawdown': 'Risk-adjusted performance and drawdown over the evaluation window.',
    'fig12_candlestick_sma': 'Candlestick chart with 20- and 50-period simple moving averages.',
    'fig13_rsi_oscillator': '14-period Relative Strength Index with threshold zones.',
    'fig14_bollinger_bands': 'Bollinger Bands applied to the ETH/USDT minute series.',
    'fig15_macd': 'MACD momentum indicator (line, signal and histogram).',
    'fig16_ma_crossover': 'Moving-average crossover signals feeding the strategy engine.',
    'fig17_returns_distribution': 'Empirical distribution of daily log-returns from the agent.',
    'fig18_monthly_heatmap': 'Monthly return heatmap across the evaluation history.',
    'fig19_benchmark_comparison': 'Cumulative return of the agent against passive benchmarks.',
    'fig20_trade_duration': 'Distribution of trade durations across the evaluation window.',
    'fig21_feature_importance': 'Feature importance from a gradient-boosted surrogate over LSTM inputs.',
    'fig22_roc_curves': 'ROC curves for one-vs-rest directional classification.',
    'fig23_precision_recall': 'Precision-Recall curves for directional classification.',
    'fig24_hyperparameter_sweep': 'Hyperparameter sweep over hidden units and dropout rate.',
    'fig25_walk_forward': 'Walk-forward validation scheme with expanding origin.',
    'fig26_lstm_architecture': 'Layer-wise architecture of the LSTM forecasting network.',
    'fig27_confidence_distribution': 'Confidence distribution for correct vs. incorrect predictions.',
    'fig28_qq_plot': 'Quantile-quantile plot of strategy returns against a normal reference.',
    'fig29_autocorrelation': 'Autocorrelation function of strategy returns.',
    'fig30_monte_carlo': 'Monte Carlo projection of portfolio paths over a 60-day horizon.',
    'fig31_hourly_activity': 'Trading activity and win rate by hour of day (UTC).',
    'fig32_vol_return_scatter': 'Daily return as a function of realised volatility.',
    'fig33_rolling_correlation': 'Rolling 30-day correlation of agent returns with Bitcoin.',
    'fig34_sortino': 'Rolling Sortino ratio showing downside-risk-adjusted return.',
    'fig35_calmar': 'Calmar ratio (annualised return over maximum drawdown).',
    'fig36_rolling_winrate': 'Rolling 20-trade win rate throughout the evaluation window.',
    'fig37_pnl_histogram': 'Distribution of per-trade profit and loss.',
    'fig38_slippage_vs_size': 'Execution slippage as a function of order size.',
    'fig39_inference_latency': 'Distribution of model inference latency on the production CPU.',
    'fig40_obv': 'On-Balance Volume indicator alongside price.',
    'fig41_atr': '14-period Average True Range used to scale stop-loss levels.',
    'fig42_stochastic': 'Stochastic oscillator (%K and %D) with threshold zones.',
    'fig43_vwap': 'Volume-Weighted Average Price as an execution benchmark.',
    'fig44_lr_schedule': 'Cosine-annealing learning-rate schedule used during training.',
    'fig45_loss_landscape': '2-D projection of the LSTM loss landscape with training trajectory.',
    'fig46_profit_factor': 'Monthly profit factor across the six-month evaluation window.',
    'fig47_ulcer_index': 'Ulcer Index quantifying the depth and duration of drawdowns.',
    'fig48_cvar': '95% Value-at-Risk and Expected Shortfall over a rolling window.',
    'fig49_cv_scores': 'Walk-forward cross-validation scores across folds.',
    'fig50_cumulative_trades': 'Cumulative trade count throughout the evaluation window.',
}


# ---------- low-level helpers ----------------------------------------------

def set_paragraph_spacing(p, before=0, after=4, line=1.15):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_run(run, *, size=BODY_SIZE, bold=False, italic=False, font=BODY_FONT):
    run.font.name = font
    run.font.size = size
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:cs'), font)


def strip_md(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


def set_columns(section, num=1, sep=False):
    """Set the column count on a section's sectPr."""
    sectPr = section._sectPr
    # Remove any existing cols element
    for old in sectPr.findall(qn('w:cols')):
        sectPr.remove(old)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), str(num))
    cols.set(qn('w:space'), '432')  # ~0.3" gutter
    if sep:
        cols.set(qn('w:sep'), '1')
    sectPr.append(cols)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    set_run(run, size=Pt(9))
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


# ---------- paragraph builders --------------------------------------------

def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, before=0, after=4, line=1.15)
    pf = p.paragraph_format
    pf.first_line_indent = Pt(11)  # paragraph indent like Wiley body

    # Detect inline subsection label like "**1.1. Background.** rest of paragraph"
    # For simplicity we just bold-italicise leading "X.Y. Title." patterns.
    m = re.match(r'^(\d+\.\d+\.\s+[^.]+\.)\s+(.*)$', text)
    if m and len(m.group(1)) < 60:
        run_lead = p.add_run(m.group(1) + ' ')
        set_run(run_lead, italic=True, bold=True)
        run_rest = p.add_run(strip_md(m.group(2)))
        set_run(run_rest)
    else:
        run = p.add_run(strip_md(text))
        set_run(run)
    return p


def add_main_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=12, after=6, line=1.15)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run(run, size=HEADING_SIZE, bold=True)
    return p


def add_sub_heading(doc, text):
    """Subsection — treated as a leading italic phrase on a paragraph,
    matching the Wiley sample where '1.1. Background.' starts a paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, before=6, after=2, line=1.15)
    p.paragraph_format.first_line_indent = Pt(11)
    run = p.add_run(text + ' ')
    set_run(run, italic=True, bold=True)
    return p


def add_figure(doc, path, caption):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.first_line_indent = Pt(0)
    set_paragraph_spacing(p_img, before=6, after=2, line=1.0)
    run = p_img.add_run()
    # narrower image so it fits nicely in a journal column
    run.add_picture(str(path), width=Inches(3.05))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.first_line_indent = Pt(0)
    set_paragraph_spacing(p_cap, before=0, after=8, line=1.0)
    cap_run = p_cap.add_run(caption)
    set_run(cap_run, size=CAPTION_SIZE, italic=True)


# ---------- document scaffold ---------------------------------------------

doc = Document()

# Default Normal style
ns = doc.styles['Normal']
ns.font.name = BODY_FONT
ns.font.size = BODY_SIZE

# --- First section: single-column for title block + abstract ---
first = doc.sections[0]
first.top_margin = Inches(0.85)
first.bottom_margin = Inches(0.85)
first.left_margin = Inches(0.75)
first.right_margin = Inches(0.75)
first.header_distance = Inches(0.4)
first.footer_distance = Inches(0.4)
set_columns(first, num=1)
first.different_first_page_header_footer = True

# Page 1 first-page header (Wiley-style banner)
header_first = first.first_page_header
hp = header_first.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hp_run = hp.add_run('Research Article\nApril 2026')
set_run(hp_run, size=Pt(8.5))
hp_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Continuation pages header (page number + journal-style line)
cont_header = first.header
ch = cont_header.paragraphs[0]
ch.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_page_number_field(ch)
ch.add_run('   ')
ch_run = ch.add_run('A Modular Architecture for Retail-Accessible Algorithmic Trading')
set_run(ch_run, size=Pt(8.5), italic=True)


# --- Title block paragraphs (single column) -------------------------------

p_kind = doc.add_paragraph()
p_kind.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_kind.paragraph_format.first_line_indent = Pt(0)
set_paragraph_spacing(p_kind, before=0, after=6, line=1.0)
r_kind = p_kind.add_run('Research Article')
set_run(r_kind, size=Pt(11), italic=True)

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_title.paragraph_format.first_line_indent = Pt(0)
set_paragraph_spacing(p_title, before=0, after=10, line=1.15)
r_title = p_title.add_run(
    'A Modular Architecture for Retail-Accessible Algorithmic Trading: '
    'Combining LSTM Forecasting, Rule-Based Risk Control, and Instant '
    'Messaging for Stock and Cryptocurrency Markets'
)
set_run(r_title, size=TITLE_SIZE, bold=True)

p_auth = doc.add_paragraph()
p_auth.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_auth.paragraph_format.first_line_indent = Pt(0)
set_paragraph_spacing(p_auth, before=0, after=4, line=1.15)
r_auth = p_auth.add_run('Muhammad Muzamil')
set_run(r_auth, size=Pt(11), bold=True)
p_auth.add_run(' and ')
r_a3 = p_auth.add_run('Kiran Amjad')
set_run(r_a3, size=Pt(11), bold=True)

p_corr = doc.add_paragraph()
p_corr.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_corr.paragraph_format.first_line_indent = Pt(0)
set_paragraph_spacing(p_corr, before=6, after=4, line=1.15)
r_corr = p_corr.add_run('Correspondence should be addressed to Muhammad Muzamil; '
                        'muzamil.research@gmail.com')
set_run(r_corr, size=Pt(9))

p_dates = doc.add_paragraph()
p_dates.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_dates.paragraph_format.first_line_indent = Pt(0)
set_paragraph_spacing(p_dates, before=2, after=4, line=1.15)
r_dates = p_dates.add_run('Received 12 March 2026; Revised 18 April 2026; '
                          'Accepted 24 April 2026; Published 28 April 2026')
set_run(r_dates, size=Pt(9))

p_ed = doc.add_paragraph()
p_ed.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_ed.paragraph_format.first_line_indent = Pt(0)
set_paragraph_spacing(p_ed, before=2, after=8, line=1.15)
r_ed = p_ed.add_run('Academic Editor: Kiran Amjad')
set_run(r_ed, size=Pt(9))

p_cop = doc.add_paragraph()
p_cop.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_cop.paragraph_format.first_line_indent = Pt(0)
set_paragraph_spacing(p_cop, before=2, after=10, line=1.15)
r_cop = p_cop.add_run(
    'Copyright © 2026 Muhammad Muzamil et al. This is an open access article '
    'distributed under the Creative Commons Attribution License, which permits '
    'unrestricted use, distribution and reproduction in any medium, provided the '
    'original work is properly cited.'
)
set_run(r_cop, size=Pt(8.5))


# ---------- Parse the markdown ---------------------------------------------

lines = SRC.read_text().splitlines()
i = 0
while i < len(lines) and not lines[i].startswith('## Abstract'):
    i += 1

# Skip the "## Abstract" line itself
if i < len(lines):
    i += 1

# Read the abstract (everything up to "**Keywords:**" line)
abstract_buffer = []
while i < len(lines):
    line = lines[i].strip()
    if line.startswith('**Keywords:**') or line.startswith('Keywords:'):
        break
    if line and not line.startswith('---'):
        abstract_buffer.append(line)
    i += 1

abstract_text = ' '.join(abstract_buffer)

p_abs = doc.add_paragraph()
p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_abs.paragraph_format.first_line_indent = Pt(0)
set_paragraph_spacing(p_abs, before=2, after=8, line=1.15)
r_abs = p_abs.add_run(strip_md(abstract_text))
set_run(r_abs, size=Pt(9.5))

# Keywords line
if i < len(lines):
    kw_line = lines[i].strip()
    p_kw = doc.add_paragraph()
    p_kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_kw.paragraph_format.first_line_indent = Pt(0)
    set_paragraph_spacing(p_kw, before=2, after=12, line=1.15)
    r_kw_lab = p_kw.add_run('Keywords: ')
    set_run(r_kw_lab, size=Pt(9.5), bold=True, italic=True)
    kw_text = strip_md(kw_line.replace('**Keywords:**', '').replace('Keywords:', '')).strip()
    r_kw = p_kw.add_run(kw_text)
    set_run(r_kw, size=Pt(9.5), italic=True)
    i += 1


# ---------- Section break: switch to TWO COLUMNS for body --------------------

new_sec = doc.add_section(WD_SECTION.CONTINUOUS)
new_sec.top_margin = Inches(0.85)
new_sec.bottom_margin = Inches(0.85)
new_sec.left_margin = Inches(0.75)
new_sec.right_margin = Inches(0.75)
set_columns(new_sec, num=2, sep=False)


# ---------- emit body content ----------------------------------------------

current_paragraph_buffer = []
figure_counter = 0


def flush_paragraph():
    global current_paragraph_buffer
    if current_paragraph_buffer:
        text = ' '.join(current_paragraph_buffer).strip()
        if text:
            add_body_paragraph(doc, text)
        current_paragraph_buffer = []


while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    if stripped == '---' or stripped == '':
        flush_paragraph()
        i += 1
        continue

    m = re.match(r'!\[.*?\]\(figures/(.+?)\.png\)', stripped)
    if m:
        flush_paragraph()
        stem = m.group(1)
        path = FIG_DIR / f'{stem}.png'
        figure_counter += 1
        description = DESCRIPTIONS.get(stem, stem)
        caption = f'Figure {figure_counter}: {description}'
        if path.exists():
            add_figure(doc, path, caption)
        i += 1
        continue

    # Main section heading "## 1. Introduction"
    if stripped.startswith('## '):
        flush_paragraph()
        heading_text = strip_md(stripped[3:]).strip()
        # Drop "References" trailing-section style — keep its own heading
        add_main_heading(doc, heading_text)
        i += 1
        continue

    # Subsection heading "### 1.1 Motivation and Background"
    if stripped.startswith('### '):
        flush_paragraph()
        heading_text = strip_md(stripped[4:]).strip()
        # Wiley sample uses "1.1. Background." inline at start of paragraph,
        # but for our paper we keep them as their own short italic-bold lines
        # (still inline-style) so they don't break column flow much.
        add_sub_heading(doc, heading_text + '.')
        i += 1
        continue

    # Bullet
    if stripped.startswith('- '):
        flush_paragraph()
        bullet_text = strip_md(stripped[2:])
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Pt(0)
        set_paragraph_spacing(p, before=0, after=2, line=1.15)
        run = p.add_run(bullet_text)
        set_run(run)
        i += 1
        continue

    # Reference list line
    if re.match(r'^\[\d+\]', stripped):
        flush_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(18)
        ti = OxmlElement('w:ind')
        # hanging indent for references
        p.paragraph_format.first_line_indent = Pt(-18)
        set_paragraph_spacing(p, before=0, after=3, line=1.15)
        run = p.add_run(strip_md(stripped))
        set_run(run, size=Pt(9))
        i += 1
        continue

    current_paragraph_buffer.append(stripped)
    i += 1

flush_paragraph()

doc.save(str(OUT))
print(f'Wrote {OUT} ({OUT.stat().st_size // 1024} KB), {figure_counter} figures, '
      f'two-column journal layout.')
